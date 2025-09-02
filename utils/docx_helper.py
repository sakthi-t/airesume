# utils/docx_helper.py
import io
from docx import Document
from utils.date_utils import format_date


def build_docx(resume: dict) -> io.BytesIO:
    """
    Build DOCX in-memory from normalized resume dict.
    Returns io.BytesIO (seeked to 0).
    """
    buffer = io.BytesIO()
    doc = Document()

    basics = resume.get("basics", {})
    name = basics.get("name", "")
    if name:
        doc.add_heading(name, level=0)

    contact_parts = [
        basics.get("email", ""),
        basics.get("phone", ""),
        basics.get("linkedin", ""),
        basics.get("github", ""),
        basics.get("website", ""),
    ]
    contact = " | ".join([p for p in contact_parts if p])
    if contact:
        doc.add_paragraph(contact)

    # Summary
    if basics.get("summary"):
        doc.add_heading("Profile Summary", level=1)
        doc.add_paragraph(basics["summary"])

    # Experience
    if resume.get("experience"):
        doc.add_heading("Work Experience", level=1)
        for e in resume["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{e.title}")
            run.bold = True
            if getattr(e, "org", None):
                p.add_run(f" — {e.org}")
            # Date range
            date_range = format_date(e.start, e.end)
            if date_range:
                p2 = doc.add_paragraph()
                r2 = p2.add_run(date_range)
                r2.italic = True
            if getattr(e, "remarks", None):
                doc.add_paragraph(e.remarks)

    # Projects
    if resume.get("projects"):
        doc.add_heading("Projects", level=1)
        for p_item in resume["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{p_item.title}")
            run.bold = True
            if getattr(p_item, "org", None):
                p.add_run(f" — {p_item.org}")
            date_range = format_date(p_item.start, p_item.end)
            if date_range:
                p2 = doc.add_paragraph()
                r2 = p2.add_run(date_range)
                r2.italic = True
            if getattr(p_item, "remarks", None):
                doc.add_paragraph(p_item.remarks)

    # Education
    if resume.get("education"):
        doc.add_heading("Education", level=1)
        for ed in resume["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{ed.title}")
            run.bold = True
            if getattr(ed, "org", None):
                p.add_run(f" — {ed.org}")
            date_range = format_date(ed.start, ed.end)
            if date_range:
                p2 = doc.add_paragraph()
                r2 = p2.add_run(date_range)
                r2.italic = True
            if getattr(ed, "remarks", None):
                doc.add_paragraph(ed.remarks)

    # Skills
    if resume.get("skills"):
        doc.add_heading("Skills", level=1)
        for s in resume["skills"]:
            doc.add_paragraph(s, style="List Bullet")

    # Awards
    if resume.get("awards"):
        doc.add_heading("Awards & Achievements", level=1)
        for a in resume["awards"]:
            doc.add_paragraph(a, style="List Bullet")

    doc.save(buffer)
    buffer.seek(0)
    return buffer
